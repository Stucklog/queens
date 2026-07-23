import json
import glob
import os
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "story-reviews"
OUT.mkdir(parents=True, exist_ok=True)


INK = "243447"
ACCENT = "2E74B5"
ACCENT_DARK = "1F4D78"
MUTED = "667085"
PALE = "E8EEF5"
PALE2 = "F4F6F9"
GREEN = "1F6654"


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths)))
    tblW.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths[i]))
            tcW.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run(run, size=11, color=INK, bold=False, italic=False, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def style_doc(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18
    for name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 18, 8),
        ("Heading 2", 13, ACCENT, 13, 6),
        ("Heading 3", 11.5, ACCENT_DARK, 9, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header.add_run("REGALIA  /  STORY CONTENT REVIEW"), 8, MUTED, True)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(footer.add_run("Reviewer working document"), 8, MUTED, False)


def p(doc, text="", style=None, before=None, after=None, align=None):
    para = doc.add_paragraph(style=style)
    if text:
        set_run(para.add_run(text))
    if before is not None:
        para.paragraph_format.space_before = Pt(before)
    if after is not None:
        para.paragraph_format.space_after = Pt(after)
    if align is not None:
        para.alignment = align
    return para


def label_value(doc, label, value):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(3)
    set_run(para.add_run(label + ": "), 10.5, INK, True)
    set_run(para.add_run(value), 10.5, INK)
    return para


def feedback_block(doc, label, lines=3, prompt=None):
    if prompt:
        q = doc.add_paragraph()
        q.paragraph_format.space_after = Pt(2)
        set_run(q.add_run(prompt), 9.5, MUTED, True)
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FBFCFD")
    cell.text = ""
    for i in range(lines):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        para.paragraph_format.space_after = Pt(5)
        set_run(para.add_run("________________________________________________________________________________"), 9.5, "98A2B3")
    cap = doc.add_paragraph()
    cap.paragraph_format.space_after = Pt(8)
    set_run(cap.add_run(label), 8.5, MUTED, True, italic=True)


def checkbox_line(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(3)
    set_run(para.add_run("☐ " + text), 10.5, INK)
    return para


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve"); instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)


def chapter_frame(arc, chapter_index):
    scenes = arc.get("scenes", [])
    if chapter_index < len(scenes):
        scene = scenes[chapter_index]
        frames = scene.get("frames", [])
        if frames:
            return frames[0]
        pages = scene.get("pages", [])
        if pages:
            return pages[0]
    return {}


def cast_from_scenes(arc):
    found = {}
    for scene in arc.get("scenes", []):
        chars = scene.get("defaults", {}).get("characters", [])
        for c in chars:
            cid = c.get("id", "")
            found[cid] = c.get("semanticLabel", cid)
    return list(found.values())


def asset(path):
    if not path:
        return "—"
    full = ROOT / path
    status = "present" if full.exists() else "MISSING"
    return f"{path} [{status}]"


def add_title_block(doc, arc):
    title = p(doc, arc["title"], after=2)
    set_run(title.runs[0], 24, INK, True)
    sub = p(doc, "Story content review worksheet", after=12)
    set_run(sub.runs[0], 12, MUTED, False, True)
    tagline = arc.get("tagline", "")
    if tagline:
        call = doc.add_table(rows=1, cols=1)
        set_table_widths(call, [9360])
        cell = call.cell(0, 0); set_cell_shading(cell, PALE2); cell.text = ""
        para = cell.paragraphs[0]; para.paragraph_format.space_after = Pt(0)
        set_run(para.add_run("STORY PROMISE  "), 8.5, ACCENT_DARK, True)
        set_run(para.add_run(tagline), 10.5, INK, False, True)
    p(doc, "", after=3)
    metrics = doc.add_table(rows=1, cols=4)
    set_table_widths(metrics, [1800, 2520, 2520, 2520])
    values = [
        ("HERO", arc.get("hero", {}).get("name", "—")),
        ("CHAPTERS", str(len(arc.get("chapters", [])))),
        ("SCENES", str(len(arc.get("scenes", [])))),
        ("REVIEW STATUS", "Not started"),
    ]
    for cell, (k, v) in zip(metrics.rows[0].cells, values):
        set_cell_shading(cell, PALE)
        cell.text = ""
        a = cell.paragraphs[0]; a.paragraph_format.space_after = Pt(2)
        set_run(a.add_run(k), 8, ACCENT_DARK, True)
        b = cell.add_paragraph(); b.paragraph_format.space_after = Pt(0)
        set_run(b.add_run(v), 10.5, INK, True)


def add_overview(doc, arc):
    doc.add_heading("How to use this worksheet", level=1)
    p(doc, "Review in order and write only the decision-relevant note at each checkpoint. Start with the story pass, then the hero and cast, then chapter backgrounds and encounter sprites, and finish with palette and system consistency. Mark anything that needs a change with a clear action verb: rewrite, remove, add, reposition, recolor, rename, or verify.")
    doc.add_heading("Fast review route", level=2)
    for text in [
        "Story pass — read the arc title, tagline, chapter titles, captions, and scene copy without looking at art.",
        "Art pass — inspect the hero, supporting characters, opponents, bosses, finishers, and backgrounds at actual display size.",
        "Consistency pass — compare palette, silhouette language, contrast, naming, and emotional escalation across chapters.",
        "Action pass — consolidate duplicate notes, rank fixes by impact, and record any final questions for implementation.",
    ]:
        checkbox_line(doc, text)
    doc.add_heading("Arc-level feedback", level=2)
    feedback_block(doc, "Overall arc feedback", 4, "What is working? What is confusing, weak, repetitive, or off-tone? What single change would improve the arc most?")


def add_story_pass(doc, arc):
    doc.add_heading("1. Story pass", level=1)
    label_value(doc, "Review target", "Narrative clarity, pacing, character motivation, escalation, payoff, and chapter-to-chapter continuity")
    label_value(doc, "Keep in mind", "You are reviewing the written experience first. Do not let a strong image hide a weak beat, or a weak image create a false story note.")
    for i, ch in enumerate(arc.get("chapters", [])):
        frame = chapter_frame(arc, i)
        doc.add_heading(f"Chapter {i+1} — {ch.get('title', 'Untitled')}", level=2)
        p(doc, ch.get("caption", "No chapter caption recorded."), after=4)
        if frame:
            label_value(doc, "Scene beat", " ".join(frame.get("paragraphs", [])))
            label_value(doc, "Intent cue", frame.get("actionLabel", "—"))
        label_value(doc, "Boss / escalation", ch.get("boss", {}).get("name", "—"))
        feedback_block(doc, f"Chapter {i+1} feedback", 2, "Is the beat clear, emotionally earned, and distinct from the chapter before it? Note exact copy or structural changes.")
    doc.add_page_break()
    doc.add_heading("Story synthesis", level=2)
    feedback_block(doc, "Story synthesis feedback", 4, "Does the arc have a clear beginning, middle, turn, climax, and ending? Are the protagonist’s choices and the final payoff legible?")


def add_art_pass(doc, arc):
    doc.add_page_break()
    doc.add_heading("2. Art pass", level=1)
    hero = arc.get("hero", {})
    label_value(doc, "Hero", hero.get("name", "—"))
    label_value(doc, "Story sprite", asset(hero.get("storySpriteAsset")))
    label_value(doc, "Combat sprite", asset(hero.get("combatSpriteAsset")))
    label_value(doc, "Finisher sprite", asset(hero.get("finisherSpriteAsset")))
    feedback_block(doc, "Hero feedback", 3, "Check silhouette, readability, pose, expression, costume/story cues, frame consistency, scale, and whether the hero belongs in this world.")
    doc.add_heading("Supporting cast and scene staging", level=2)
    cast = cast_from_scenes(arc)
    if cast:
        for item in cast:
            p(doc, item, style="List Bullet", after=3)
    else:
        p(doc, "No supporting cast entries were found in the current scene defaults.")
    feedback_block(doc, "Cast and staging feedback", 3, "Check whether every visible character has a clear role, readable placement, consistent scale, and a useful relationship to the background and text.")
    doc.add_heading("Combat sprite inventory", level=2)
    for i, ch in enumerate(arc.get("chapters", [])):
        p(doc, f"Chapter {i+1} — {ch.get('title', 'Untitled')}", style="Heading 3", after=2)
        boss = ch.get("boss", {})
        label_value(doc, "Boss", f"{boss.get('name', '—')}  |  {asset(boss.get('spriteAsset'))}")
        for enemy in ch.get("encounters", []):
            label_value(doc, "Encounter", f"{enemy.get('name', '—')}  |  {asset(enemy.get('spriteAsset'))}")
    doc.add_heading("Chapter art inventory", level=2)
    table = doc.add_table(rows=1, cols=5)
    set_table_widths(table, [650, 1900, 2150, 2350, 2310])
    headers = ["#", "Chapter", "Background", "Boss / opponent sprites", "Palette"]
    for cell, h in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, PALE); cell.text = ""; set_run(cell.paragraphs[0].add_run(h), 8.5, ACCENT_DARK, True)
    for i, ch in enumerate(arc.get("chapters", [])):
        row = table.add_row().cells
        vals = [str(i+1), ch.get("title", "—"), asset(ch.get("artAsset")), ch.get("boss", {}).get("name", "—") + "; " + ", ".join(e.get("name", "—") for e in ch.get("encounters", [])), ch.get("primaryColor", "—") + " / " + ch.get("secondaryColor", "—")]
        for cell, v in zip(row, vals):
            cell.text = ""; set_run(cell.paragraphs[0].add_run(v), 8.5, INK)
    feedback_block(doc, "Art inventory feedback", 3, "For each asset above, record what must change and why. Flag missing files, mismatched names, duplicate silhouettes, weak chapter progression, or poor crop/contrast.")


def add_system_pass(doc, arc):
    doc.add_heading("3. Color, consistency, and finish", level=1)
    theme = arc.get("theme", {})
    label_value(doc, "Brightness", theme.get("brightness", "—"))
    label_value(doc, "Background", theme.get("backgroundColor", "—"))
    label_value(doc, "Surface", theme.get("surfaceColor", "—"))
    label_value(doc, "Foreground / ink", theme.get("foregroundColor", "—") + " / " + theme.get("inkColor", "—"))
    label_value(doc, "Danger", theme.get("dangerColor", "—"))
    p(doc, "Compare the theme colors against the chapter accent pairs above. Check readability in both quiet story scenes and high-intensity combat moments; the palette should support hierarchy, not just atmosphere.")
    for text in [
        "Palette — colors are intentional, distinguishable, accessible, and consistent with the story’s emotional arc.",
        "Backgrounds — focal areas remain clear behind characters and copy; no important detail is lost to crop or contrast.",
        "Sprites — silhouettes, edge treatment, animation frames, and scale feel like one production family.",
        "Naming — titles, captions, IDs, filenames, and visible labels agree.",
        "Escalation — chapter art and boss spectacle build toward the finale rather than peaking too early.",
        "Polish — no accidental artifacts, broken alpha, awkward tangents, duplicated motifs, or unfinished frames.",
    ]:
        checkbox_line(doc, text)
    feedback_block(doc, "Color and consistency feedback", 4, "Write specific corrections, not just preferences. Name the asset and the proposed direction when possible.")
    doc.add_heading("4. Final action list", level=1)
    table = doc.add_table(rows=1, cols=4)
    set_table_widths(table, [1750, 4100, 2000, 1510])
    for cell, h in zip(table.rows[0].cells, ["Priority", "Correction", "Asset / location", "Status"]):
        set_cell_shading(cell, PALE); cell.text = ""; set_run(cell.paragraphs[0].add_run(h), 9, ACCENT_DARK, True)
    for _ in range(6):
        row = table.add_row().cells
        for cell, v in zip(row, ["P0 / P1 / P2", "", "", "Open"]):
            cell.text = ""; set_run(cell.paragraphs[0].add_run(v), 9, "98A2B3", False, v == "")
    p(doc, "", after=2)
    feedback_block(doc, "Reviewer sign-off", 2, "When finished: note any unresolved questions, the highest-impact correction, and whether the story is ready for implementation review.")


def md_feedback(label, prompt, lines=3):
    out = [f"**{prompt}**", "", f"_{label}_", ""]
    out.extend(["- "] * lines)
    out.append("")
    return out


def write_markdown(arc, slug, spec_path):
    lines = [f"# {arc['title']}", "", "_Story content review worksheet_", ""]
    if arc.get("tagline"):
        lines += [f"> **Story promise:** {arc['tagline']}", ""]
    hero = arc.get("hero", {}).get("name", "—")
    lines += [
        "| Hero | Chapters | Scenes | Review status |",
        "|---|---:|---:|---|",
        f"| {hero} | {len(arc.get('chapters', []))} | {len(arc.get('scenes', []))} | Not started |",
        "",
        f"**Source files:** {arc['story_spec_source']}" + (f" + assets/content/arcs/{slug}/arc.json" if spec_path else ""),
        "",
        "## How to use this worksheet",
        "",
        "Review in order and write only the decision-relevant note at each checkpoint. Start with the story pass, then the hero and cast, then chapter backgrounds and encounter sprites, and finish with palette and system consistency. Mark anything that needs a change with a clear action verb: rewrite, remove, add, reposition, recolor, rename, or verify.",
        "",
        "### Fast review route",
        "",
        "- [ ] Story pass — read the arc title, tagline, chapter titles, captions, and scene copy without looking at art.",
        "- [ ] Art pass — inspect the hero, supporting characters, opponents, bosses, finishers, and backgrounds at actual display size.",
        "- [ ] Consistency pass — compare palette, silhouette language, contrast, naming, and emotional escalation across chapters.",
        "- [ ] Action pass — consolidate duplicate notes, rank fixes by impact, and record any final questions for implementation.",
        "",
        "### Arc-level feedback",
        "",
    ]
    lines += md_feedback("Overall arc feedback", "What is working? What is confusing, weak, repetitive, or off-tone? What single change would improve the arc most?", 4)
    lines += ["## 1. Story pass", "", "**Review target:** Narrative clarity, pacing, character motivation, escalation, payoff, and chapter-to-chapter continuity", "", "**Keep in mind:** You are reviewing the written experience first. Do not let a strong image hide a weak beat, or a weak image create a false story note.", ""]
    for i, ch in enumerate(arc.get("chapters", [])):
        frame = chapter_frame(arc, i)
        lines += [f"### Chapter {i+1} — {ch.get('title', 'Untitled')}", "", ch.get("caption", "No chapter caption recorded."), ""]
        if frame:
            lines += [f"**Scene beat:** {' '.join(frame.get('paragraphs', []))}", "", f"**Intent cue:** {frame.get('actionLabel', '—')}", ""]
        lines += [f"**Boss / escalation:** {ch.get('boss', {}).get('name', '—')}", ""]
        lines += md_feedback(f"Chapter {i+1} feedback", "Is the beat clear, emotionally earned, and distinct from the chapter before it? Note exact copy or structural changes.", 2)
    lines += ["## Story synthesis", ""]
    lines += md_feedback("Story synthesis feedback", "Does the arc have a clear beginning, middle, turn, climax, and ending? Are the protagonist’s choices and the final payoff legible?", 4)
    lines += ["## 2. Art pass", "", f"**Hero:** {hero}"]
    h = arc.get("hero", {})
    for label, key in [("Story sprite", "storySpriteAsset"), ("Combat sprite", "combatSpriteAsset"), ("Finisher sprite", "finisherSpriteAsset")]:
        if h.get(key): lines.append(f"- **{label}:** `{h[key]}`")
    lines.append("")
    lines += md_feedback("Hero feedback", "Check silhouette, readability, pose, expression, costume/story cues, frame consistency, scale, and whether the hero belongs in this world.", 3)
    lines += ["### Supporting cast and scene staging", ""]
    cast = cast_from_scenes(arc)
    lines += [f"- {x}" for x in cast] if cast else ["No supporting cast entries were found in the current scene defaults."]
    lines.append("")
    lines += md_feedback("Cast and staging feedback", "Check whether every visible character has a clear role, readable placement, consistent scale, and a useful relationship to the background and text.", 3)
    lines += ["### Combat sprite inventory", ""]
    for i, ch in enumerate(arc.get("chapters", [])):
        lines += [f"#### Chapter {i+1} — {ch.get('title', 'Untitled')}", ""]
        boss = ch.get("boss", {})
        lines.append(f"- **Boss:** {boss.get('name', '—')} — `{boss.get('spriteAsset', '—')}`")
        for enemy in ch.get("encounters", []):
            lines.append(f"- **Encounter:** {enemy.get('name', '—')} — `{enemy.get('spriteAsset', '—')}`")
        lines.append("")
    lines += ["### Chapter art inventory", "", "| # | Chapter | Background | Boss / opponent sprites | Palette |", "|---:|---|---|---|---|"]
    for i, ch in enumerate(arc.get("chapters", [])):
        names = ch.get('boss', {}).get('name', '—') + "; " + ", ".join(e.get('name', '—') for e in ch.get('encounters', []))
        palette = ch.get('primaryColor', '—') + " / " + ch.get('secondaryColor', '—')
        lines.append(f"| {i+1} | {ch.get('title', '—')} | `{ch.get('artAsset', '—')}` | {names} | `{palette}` |")
    lines.append("")
    lines += md_feedback("Art inventory feedback", "For each asset above, record what must change and why. Flag missing files, mismatched names, duplicate silhouettes, weak chapter progression, or poor crop/contrast.", 3)
    theme = arc.get("theme", {})
    lines += ["## 3. Color, consistency, and finish", "", f"- **Brightness:** {theme.get('brightness', '—')}", f"- **Background:** `{theme.get('backgroundColor', '—')}`", f"- **Surface:** `{theme.get('surfaceColor', '—')}`", f"- **Foreground / ink:** `{theme.get('foregroundColor', '—')} / {theme.get('inkColor', '—')}`", f"- **Danger:** `{theme.get('dangerColor', '—')}`", "", "Compare the theme colors against the chapter accent pairs above. Check readability in both quiet story scenes and high-intensity combat moments; the palette should support hierarchy, not just atmosphere.", "", "- [ ] Palette — colors are intentional, distinguishable, accessible, and consistent with the story’s emotional arc.", "- [ ] Backgrounds — focal areas remain clear behind characters and copy; no important detail is lost to crop or contrast.", "- [ ] Sprites — silhouettes, edge treatment, animation frames, and scale feel like one production family.", "- [ ] Naming — titles, captions, IDs, filenames, and visible labels agree.", "- [ ] Escalation — chapter art and boss spectacle build toward the finale rather than peaking too early.", "- [ ] Polish — no accidental artifacts, broken alpha, awkward tangents, duplicated motifs, or unfinished frames.", ""]
    lines += md_feedback("Color and consistency feedback", "Write specific corrections, not just preferences. Name the asset and the proposed direction when possible.", 4)
    lines += ["## 4. Final action list", "", "| Priority | Correction | Asset / location | Status |", "|---|---|---|---|"]
    lines += ["| P0 / P1 / P2 |  |  | Open |" for _ in range(6)]
    lines += ["", "**When finished:** note any unresolved questions, the highest-impact correction, and whether the story is ready for implementation review.", ""]
    lines += md_feedback("Reviewer sign-off", "", 2)
    (OUT / f"{slug}-story-review.md").write_text("\n".join(lines) + "\n")


def main():
    specs = {Path(p).stem: Path(p) for p in glob.glob(str(ROOT / "tool" / "story_arc_specs" / "*.json"))}
    slugs = sorted(set(specs) | {"origin"})
    for slug in slugs:
        spec_path = specs.get(slug)
        content_path = ROOT / "assets" / "content" / "arcs" / slug / "arc.json"
        if not content_path.exists():
            continue
        arc = json.loads(content_path.read_text())
        spec = json.loads(spec_path.read_text()) if spec_path else {}
        arc["tagline"] = spec.get("tagline", "")
        arc["story_spec_source"] = str(spec_path.relative_to(ROOT)) if spec_path else f"assets/content/arcs/{slug}/arc.json"
        out = OUT / f"{slug}-story-review.docx"
        doc = Document()
        style_doc(doc)
        add_title_block(doc, arc)
        if spec_path:
            source_files = f"{arc['story_spec_source']} + assets/content/arcs/{slug}/arc.json"
        else:
            source_files = arc['story_spec_source']
        label_value(doc, "Source files", source_files)
        add_overview(doc, arc)
        add_story_pass(doc, arc)
        add_art_pass(doc, arc)
        add_system_pass(doc, arc)
        # page numbering in footer
        footer = doc.sections[0].footer.paragraphs[0]
        set_run(footer.add_run("  |  Page "), 8, MUTED)
        add_page_number(footer)
        doc.save(out)
        write_markdown(arc, slug, spec_path)
        print(out)


if __name__ == "__main__":
    main()
